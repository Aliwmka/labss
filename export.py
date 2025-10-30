# export.py (исправленный)
from docx import Document
import pandas as pd
import os
from datetime import datetime

def export_to_word(self):
    try:
        rows = [(self.results_table.item(item)["values"]) for item in self.results_table.get_children()]
        
        if not rows:
            from tkinter import messagebox
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта!")
            return
        
        doc = Document()
        doc.add_heading('Отчет по бронированиям гостиницы', 0)
        
        # Добавляем информацию о фильтрах
        client_option = self.client_combobox.get() or "все"
        room_option = self.room_combobox.get() or "все"
        date_from_option = self.date_from_entry.get() or "все"
        date_to_option = self.date_to_entry.get() or "все"
        
        filters_info = f"Фильтры: Клиент - {client_option}, Номер - {room_option}, Дата заселения: {date_from_option} - {date_to_option}"
        doc.add_paragraph(filters_info)
        doc.add_paragraph()  # Пустая строка
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=len(self.results_table["columns"]))
        table.style = 'Table Grid'
        
        # Заголовки
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(self.results_table["columns"]):
            hdr_cells[i].text = str(column)
        
        # Данные
        for row in rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if value is not None else ""
        
        # Создаем папку для отчетов, если ее нет
        if not os.path.exists("reports"):
            os.makedirs("reports")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"reports/Отчет_бронирования_{timestamp}.docx"
        
        doc.save(filename)
        
        from tkinter import messagebox
        messagebox.showinfo("Успех", f"Отчет сохранен в {filename}")
        
    except Exception as e:
        from tkinter import messagebox
        messagebox.showerror("Ошибка", f"Не удалось экспортировать в Word: {e}")

def export_to_excel(self):
    try:
        rows = [(self.results_table.item(item)["values"]) for item in self.results_table.get_children()]
        
        if not rows:
            from tkinter import messagebox
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта!")
            return
        
        columns = self.results_table["columns"]
        
        # Создаем DataFrame
        df = pd.DataFrame(rows, columns=columns)
        
        # Добавляем информацию о фильтрах
        client_option = self.client_combobox.get() or "все"
        room_option = self.room_combobox.get() or "все"
        date_from_option = self.date_from_entry.get() or "все"
        date_to_option = self.date_to_entry.get() or "все"
        
        # Создаем папку для отчетов, если ее нет
        if not os.path.exists("reports"):
            os.makedirs("reports")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"reports/Отчет_бронирования_{timestamp}.xlsx"
        
        # Сохраняем в Excel с настройками
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Бронирования', index=False)
            
            # Настройка ширины колонок
            worksheet = writer.sheets['Бронирования']
            for idx, col in enumerate(df.columns):
                max_len = max(df[col].astype(str).str.len().max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_len, 50)
        
        from tkinter import messagebox
        messagebox.showinfo("Успех", f"Отчет сохранен в {filename}")
        
    except Exception as e:
        from tkinter import messagebox
        messagebox.showerror("Ошибка", f"Не удалось экспортировать в Excel: {e}")