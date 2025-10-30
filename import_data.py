# import_data.py
from docx import Document
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox

def import_from_excel(parent_window, callback):
    filepath = filedialog.askopenfilename(
        parent=parent_window,
        title="Выберите файл Excel",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    if not filepath:
        return

    try:
        # Читаем Excel файл
        df = pd.read_excel(filepath, dtype=str)
        
        # Заменяем NaN на пустые строки
        df = df.fillna('')
        
        # Преобразуем в список
        data = df.values.tolist()
        headers = df.columns.tolist()
        
        # Проверяем структуру данных
        if len(data) == 0:
            messagebox.showwarning("Предупреждение", "Файл не содержит данных!")
            return
            
        # Проверяем количество столбцов
        if len(headers) < 5:
            messagebox.showwarning("Предупреждение", 
                                "Файл должен содержать как минимум 5 столбцов: "
                                "Клиент, Номер, Дата заселения, Дата выселения, Примечание")
            return
        
        callback(headers, data)
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось прочитать Excel файл: {str(e)}")

def import_from_word(parent_window, callback):
    filepath = filedialog.askopenfilename(
        parent=parent_window,
        title="Выберите файл Word",
        filetypes=[("Word files", "*.docx")]
    )
    if not filepath:
        return

    try:
        doc = Document(filepath)
        if not doc.tables:
            messagebox.showwarning("Внимание", "В документе нет таблиц.")
            return

        table = doc.tables[0]
        data = []
        headers = []
        
        # Читаем заголовки из первой строки
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                row_data.append(text)
            
            if i == 0:
                headers = row_data
                # Проверяем количество столбцов в заголовках
                if len(headers) < 5:
                    messagebox.showwarning("Предупреждение", 
                                        "Таблица должна содержать как минимум 5 столбцов: "
                                        "Клиент, Номер, Дата заселения, Дата выселения, Примечание")
                    return
            else:
                # Пропускаем пустые строки
                if any(cell for cell in row_data):
                    # Если в данных строке меньше столбцов чем в заголовках, дополняем пустыми значениями
                    while len(row_data) < len(headers):
                        row_data.append('')
                    data.append(row_data)
        
        if not data:
            messagebox.showwarning("Внимание", "В таблице нет данных (кроме заголовков).")
            return
            
        callback(headers, data)
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось прочитать Word документ: {str(e)}")

def import_clients_from_excel(parent_window, callback):
    """Импорт клиентов из Excel"""
    filepath = filedialog.askopenfilename(
        parent=parent_window,
        title="Выберите файл Excel с клиентами",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    if not filepath:
        return

    try:
        df = pd.read_excel(filepath, dtype=str)
        df = df.fillna('')
        data = df.values.tolist()
        headers = df.columns.tolist()
        
        if len(data) == 0:
            messagebox.showwarning("Предупреждение", "Файл не содержит данных!")
            return
            
        if len(headers) < 5:
            messagebox.showwarning("Предупреждение", 
                                "Файл должен содержать как минимум 5 столбцов: "
                                "Фамилия, Имя, Отчество, Паспортные данные, Комментарий")
            return
        
        callback(headers, data)
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось прочитать Excel файл: {str(e)}")

def import_rooms_from_excel(parent_window, callback):
    """Импорт номеров из Excel"""
    filepath = filedialog.askopenfilename(
        parent=parent_window,
        title="Выберите файл Excel с номерами",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    if not filepath:
        return

    try:
        df = pd.read_excel(filepath, dtype=str)
        df = df.fillna('')
        data = df.values.tolist()
        headers = df.columns.tolist()
        
        if len(data) == 0:
            messagebox.showwarning("Предупреждение", "Файл не содержит данных!")
            return
            
        if len(headers) < 4:
            messagebox.showwarning("Предупреждение", 
                                "Файл должен содержать как минимум 4 столбца: "
                                "Номер, Вместимость, Комфортность, Цена")
            return
        
        callback(headers, data)
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось прочитать Excel файл: {str(e)}")